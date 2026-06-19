from docx import Document
from datetime import datetime
from typing import Dict, Union


def _format_analysis(analysis_result: Dict) -> str:
    lines = [
        f"Offence: {analysis_result.get('offence', '')}",
        f"Overall strength: {analysis_result.get('overall_strength', '')}",
        "",
        "Element-by-element assessment:",
    ]
    for el in analysis_result.get("elements_analysis", []):
        lines.append(
            f"- {el.get('element', '')}: {el.get('strength', '').upper()} — {el.get('details', '')}"
        )

    recommendations = analysis_result.get("recommendations", [])
    if recommendations:
        lines.append("")
        lines.append("Recommendations:")
        for rec in recommendations:
            lines.append(f"- {rec}")

    return "\n".join(lines)


def _format_law(law_result: Dict) -> str:
    lines = [
        f"Offence: {law_result.get('offence', '')}",
        "",
        "Applicable statutes:",
    ]
    for statute in law_result.get("statutes", []):
        lines.append(f"- {statute}")

    lines.append("")
    lines.append("Statutory elements:")
    for element in law_result.get("elements", []):
        lines.append(f"- {element}")

    matched = law_result.get("matched_elements", [])
    if matched:
        lines.append("")
        lines.append("Matched elements (supported by evidence):")
        for element in matched:
            lines.append(f"- {element}")

    unmatched = law_result.get("unmatched_elements", [])
    if unmatched:
        lines.append("")
        lines.append("Missing / unmatched elements:")
        for element in unmatched:
            lines.append(f"- {element}")

    checks = law_result.get("procedural_checks", [])
    if checks:
        lines.append("")
        lines.append("Procedural checks:")
        for check in checks:
            lines.append(f"- {check}")

    notes = law_result.get("notes", "")
    if notes:
        lines.append("")
        lines.append(f"Note: {notes}")

    return "\n".join(lines)


def _extract_cross(cross: Union[str, Dict]) -> str:
    if isinstance(cross, dict):
        return cross.get("cross_examination", "")
    return cross or ""


def generate_case_report(
    evidence_summary: str,
    analysis_result: Dict,
    law_result: Dict,
    cross_examination: Union[str, Dict],
) -> str:
    doc = Document()

    doc.add_heading("TTPS Case Analysis Report", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("1. Evidence Summary", level=2)
    doc.add_paragraph(evidence_summary)

    doc.add_heading("2. Applicable Law & Statutes", level=2)
    doc.add_paragraph(_format_law(law_result))

    doc.add_heading("3. Case Analysis — Strengths & Weaknesses", level=2)
    doc.add_paragraph(_format_analysis(analysis_result))

    doc.add_heading("4. King's Counsel Style Cross-Examination", level=2)
    doc.add_paragraph(_extract_cross(cross_examination))

    filename = "case_report.docx"
    doc.save(filename)
    return filename
